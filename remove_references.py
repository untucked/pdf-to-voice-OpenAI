from docx import Document
import re
import os

def clean_docx_references(input_path):
    """
    Removes in-text references like [1], [12], etc. from a .docx file,
    stopping at the 'References' section. Saves the cleaned file in the
    same directory with '_cleaned' appended to the filename.
    """
    # Compile reference pattern ([1], [12], [1][2][3], etc.)
    reference_pattern = re.compile(r"(\[\d+\])+")
    
    # Load document
    doc = Document(input_path)
    
    # Create new document
    cleaned_doc = Document()
    
    # Flag to skip after References section
    in_references_section = False
    
    for para in doc.paragraphs:
        if para.text.strip().lower().startswith("references"):
            in_references_section = True
        
        if not in_references_section:
            cleaned_text = reference_pattern.sub("", para.text)
            cleaned_doc.add_paragraph(cleaned_text)
    
    # Create output path
    base, ext = os.path.splitext(input_path)
    output_path = f"{base}_cleaned{ext}"
    
    # Save cleaned document
    cleaned_doc.save(output_path)
    
    return output_path


if __name__ == "__main__":
    # Example path – replace with your own
    input_path = "/mnt/data/General Motors (GM) Stock Deep Dive Report.docx"
    new_file = clean_docx_references(input_path)
    print(f"✅ Cleaned file saved to: {new_file}")
